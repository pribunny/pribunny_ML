from docx import Document
import re
import os

# 파일 경로 설정
input_dir = r"경로 설정 알아서"
output_dir = r"경로 설정 알아서"
os.makedirs(output_dir, exist_ok=True)

# 파일 불러오기
filename = next((f for f in os.listdir(input_dir) if f.endswith(".docx")), None)
if not filename:
    raise FileNotFoundError("DOCX 파일을 찾을 수 없습니다.")

doc = Document(os.path.join(input_dir, filename))

def paragraph_has_bold(para):
    return any(run.bold for run in para.runs if run.text.strip())

# 삭제 감지: 제목만 기준
삭제제목패턴 = re.compile(r"삭제", re.IGNORECASE)

# 추적용 변수
last_chapter = ""
last_section = ""  # 장이 바뀌면 절 초기화 필요
article_starts = []
chapter_starts = []
section_starts = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue

    # 장 감지
    m_chap = re.match(r"(제\d+장)\s+(.+)", text)
    if m_chap:
        chapter_title = m_chap.group(2).strip()
        last_chapter = f"{m_chap.group(1)} {chapter_title}"
        last_section = ""  # 🔹 새로운 장이 시작되면 절 초기화
        chapter_starts.append(i)
        continue

    # 절 감지
    m_sect = re.match(r"(제\d+절)\s*\(?([^)]+)?\)?", text)
    if m_sect:
        section_title = (m_sect.group(2) or "").strip()
        last_section = f"{m_sect.group(1)}({section_title})"
        section_starts.append(i)
        continue

    # 조 감지 (굵은 글씨)
    if paragraph_has_bold(para):
        m_article = re.match(r"(제\d+조(?:의\d+)?)(?:\s*\(([^)]+)\)|\s+(.+))?", text)
        if m_article:
            조번호 = m_article.group(1)
            제목 = m_article.group(2) or m_article.group(3) or ""
            article_starts.append((i, last_chapter, last_section, 조번호.strip(), 제목.strip(), text))

# chunking 및 저장
for idx, (para_idx, 장, 절, 조, 제목, 첫문단) in enumerate(article_starts):
    start = para_idx + 1

    # 다음 조항, 장, 절의 시작 위치 중 가장 작은 값을 선택하여 종료 위치로 설정
    end_candidates = [article_starts[idx + 1][0] if idx + 1 < len(article_starts) else len(doc.paragraphs)]
    end_candidates += [pos for pos in chapter_starts if pos > para_idx]
    end_candidates += [pos for pos in section_starts if pos > para_idx]
    end = min(end_candidates)

    나머지 = "\n".join(p.text.strip() for p in doc.paragraphs[start:end] if p.text.strip())
    full_text = 첫문단 + "\n" + 나머지

    # 삭제 여부 판단
    삭제문서 = any([
        삭제제목패턴.search(장),
        삭제제목패턴.search(절),
        삭제제목패턴.search(조),
        삭제제목패턴.search(제목)
    ])
    삭제_suffix = "삭제" if 삭제문서 else ""

    def sanitize_filename(filename):
        filename = re.sub(r'<[^>]*>', '', filename)  # 🔹 < > 안의 내용 삭제
        filename = re.sub(r'[\\/:*?"<>|]', '_', filename)  # 🔹 Windows 금지 문자 제거
        return filename.strip('_')  # 앞뒤 불필요한 _ 제거

    # 🔹 절이 없는 경우 "미지정절"로 저장
    장_part = 장 if 장 else "미지정장"
    절_part = 절 if 절 else "미지정절"
    조_part = f"{조}({제목})" if 제목 else 조
    파일명 = f"{장_part}_{절_part}_{조_part}_{삭제_suffix}.txt"
    파일명 = sanitize_filename(파일명)
    파일경로 = os.path.join(output_dir, 파일명)

    # 저장
    with open(파일경로, "w", encoding="utf-8") as f:
        f.write(full_text)

print(f"처리 완료: {len(article_starts)}개 조항 저장됨.")
